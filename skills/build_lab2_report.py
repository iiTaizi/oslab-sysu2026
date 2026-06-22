from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("2") / "操作系统实验报告2_四项任务填写版.docx"
FONT_CN = "宋体"
FONT_CODE = "Consolas"
ACCENT = "1F4E79"
MUTED = "666666"
PLACEHOLDER_FILL = "F2F2F2"
PLACEHOLDER_BORDER = "A6A6A6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name=FONT_CN, size=12, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(paragraph, first_line=True, after=4, line=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    fmt = paragraph.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(24)
    else:
        fmt.first_line_indent = Pt(0)


def add_body(doc, text, first_line=True, after=4, bold=False, color=None):
    p = doc.add_paragraph()
    set_para_format(p, first_line=first_line, after=after)
    set_run_font(p.add_run(text), bold=bold, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(10)
    fmt.space_after = Pt(6)
    fmt.keep_with_next = True
    set_run_font(p.add_run(text), size=16, bold=True, color=ACCENT)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(4)
    fmt.keep_with_next = True
    set_run_font(p.add_run(text), size=14, bold=True, color=ACCENT)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(5)
    fmt.space_after = Pt(3)
    fmt.keep_with_next = True
    set_run_font(p.add_run(text), size=12, bold=True)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para_format(p, first_line=False, after=2)
    set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_format(p, first_line=False, after=2)
    set_run_font(p.add_run(text))
    return p


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    set_cell_shading(cell, "EAF2F8")
    set_cell_border(cell, top={"val": "single", "sz": 6, "color": "9CC2E5"},
                    bottom={"val": "single", "sz": 6, "color": "9CC2E5"},
                    left={"val": "single", "sz": 6, "color": "9CC2E5"},
                    right={"val": "single", "sz": 6, "color": "9CC2E5"})
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    set_para_format(p, first_line=False, after=0)
    set_run_font(p.add_run(text), size=11, color="1F1F1F")
    doc.add_paragraph()


def add_figure_placeholder(doc, caption, prompt):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(15.6)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, PLACEHOLDER_FILL)
    set_cell_border(cell, top={"val": "dashed", "sz": 8, "color": PLACEHOLDER_BORDER},
                    bottom={"val": "dashed", "sz": 8, "color": PLACEHOLDER_BORDER},
                    left={"val": "dashed", "sz": 8, "color": PLACEHOLDER_BORDER},
                    right={"val": "dashed", "sz": 8, "color": PLACEHOLDER_BORDER})
    set_cell_margins(cell, top=500, bottom=500)
    p = cell.paragraphs[0]
    set_para_format(p, first_line=False, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.add_run(f"[截图占位]\n{prompt}"), size=11, bold=True, color=MUTED)
    cap = doc.add_paragraph()
    set_para_format(cap, first_line=False, after=5, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(cap.add_run(caption), size=10.5, color=MUTED)


def add_code_placeholder(doc, title, prompt):
    add_h3(doc, title)
    add_body(doc, prompt, first_line=False, after=2)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(15.6)
    set_cell_shading(cell, "FAFAFA")
    set_cell_border(cell, top={"val": "single", "sz": 6, "color": "BFBFBF"},
                    bottom={"val": "single", "sz": 6, "color": "BFBFBF"},
                    left={"val": "single", "sz": 6, "color": "BFBFBF"},
                    right={"val": "single", "sz": 6, "color": "BFBFBF"})
    set_cell_margins(cell, top=180, bottom=420)
    p = cell.paragraphs[0]
    set_para_format(p, first_line=False, after=0, line=1.0)
    set_run_font(p.add_run("[此处粘贴关键代码]\n\n\n"), name=FONT_CODE, size=10.5, color=MUTED)
    doc.add_paragraph()


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("操作系统原理实验报告"), size=24, bold=True)
    for _ in range(5):
        doc.add_paragraph()
    for label, value in (
        ("实验名称", "实验二  实模式和保护模式下的 OS 启动"),
        ("授课教师", "张青"),
        ("学生姓名", "________________"),
        ("学生学号", "________________"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(4.2)
        p.paragraph_format.space_after = Pt(8)
        set_run_font(p.add_run(f"{label}：  {value}"), size=14)
    doc.add_page_break()


def add_requirements(doc):
    add_h1(doc, "1. 实验要求")
    add_body(doc, "本实验围绕 IA-32 处理器、x86 汇编、实模式中断和计算机开机启动过程展开。报告主要记录以下四项任务的完成过程。")
    items = [
        "Assignment 1：复现教程中的 MBR 示例，并修改 MBR，使其加载到 0x7C00 后从屏幕坐标 (12,12) 开始输出本人学号；显示使用与教程不同的前景色和背景色。",
        "Assignment 2：探索实模式下的光标中断和键盘中断；实现光标位置获取、光标移动、通过中断输出学号，以及键盘输入并回显。",
        "Assignment 3：使用 32 位寄存器完成分支逻辑、循环逻辑和字符数组遍历函数，在 assignment 目录下通过 make run 验证。",
        "Assignment 4：编写不超过 510 字节的字符弹射程序，使字符从坐标 (2,0) 出发，以 45 度方向运动并在边界处反弹；可按个人实现加入变色或双向射出等扩展效果。",
    ]
    for item in items:
        add_numbered(doc, item)
    add_note(doc, "填写提示：本版本只依据实验 PDF 整理任务结构。请在完成实验后补充实际操作、截图、关键代码和排错记录；不要保留与本人实现不一致的描述。")


def add_process(doc):
    add_h1(doc, "2. 实验过程")
    add_body(doc, "以下内容按照四个 assignment 的顺序组织。每一部分均预留了后续补充位置，请根据本人实际实现填写命令、文件名、修改思路和调试过程。")

    add_h2(doc, "2.1 Assignment 1：MBR 基础与显存输出")
    add_h3(doc, "2.1.1 复现 Example 1")
    add_body(doc, "首先依据实验材料复现 Example 1，完成 MBR 汇编、写入镜像和启动验证。该步骤用于确认 16 位实模式程序能够被 BIOS 加载到 0x7C00 并正常执行。")
    add_body(doc, "[此处补充实际使用的汇编、写入镜像和启动命令，以及本人遇到的问题。]", color=MUTED)
    add_h3(doc, "2.1.2 在指定位置输出学号")
    add_body(doc, "在 Example 1 的基础上修改显示逻辑，使程序从屏幕坐标 (12,12) 开始输出本人学号，并调整字符属性，使前景色和背景色与教程示例不同。")
    add_body(doc, "[此处补充本人选择的颜色、坐标计算方法及修改步骤。]", color=MUTED)

    add_h2(doc, "2.2 Assignment 2：实模式中断")
    add_h3(doc, "2.2.1 光标位置获取与移动")
    add_body(doc, "通过 BIOS 视频服务中断探索光标相关功能：先获取当前光标位置，再修改寄存器参数调用中断，将光标移动到指定位置。")
    add_body(doc, "[此处补充所使用的中断号、功能号、寄存器参数和验证方法。]", color=MUTED)
    add_h3(doc, "2.2.2 使用中断输出学号")
    add_body(doc, "将 Assignment 1 中直接操作显示位置的实现改写为基于实模式中断的输出方式，逐字符显示本人学号。")
    add_body(doc, "[此处补充输出字符时使用的中断功能、循环设计和颜色设置。]", color=MUTED)
    add_h3(doc, "2.2.3 键盘输入与回显")
    add_body(doc, "继续探索 BIOS 键盘中断，读取键盘输入并将获得的字符回显到屏幕。")
    add_body(doc, "[此处补充键盘中断调用方式、结束条件、特殊按键处理和实际调试过程。]", color=MUTED)

    add_h2(doc, "2.3 Assignment 3：32 位汇编逻辑")
    add_body(doc, "该任务在 assignment 目录下完成，不需要写入 MBR 或通过 QEMU 启动。根据题目要求，在 student.asm 中分别实现分支逻辑、循环逻辑和字符数组遍历函数，并通过 make run 执行测试。")
    add_h3(doc, "2.3.1 分支逻辑")
    add_body(doc, "根据 a1 的取值范围分别计算 if_flag：当 a1 < 12 时计算 a1 / 2 + 1；当 12 <= a1 < 24 时计算 (24 - a1) * a1；否则计算 a1 << 4。")
    add_h3(doc, "2.3.2 循环逻辑")
    add_body(doc, "当 a2 >= 12 时循环调用 my_random，并将返回值依次写入 while_flag[a2 - 12]，随后递减 a2，直到循环结束。")
    add_h3(doc, "2.3.3 函数实现")
    add_body(doc, "实现 your_function，遍历以空字符结尾的字符串，并逐字符调用 print_a_char 输出。")
    add_body(doc, "[此处补充本人修改 student.asm 和测试参数的过程，以及 make run 的实际输出说明。]", color=MUTED)

    add_h2(doc, "2.4 Assignment 4：字符弹射程序")
    add_body(doc, "编写可放入 MBR 的字符弹射程序。字符从坐标 (2,0) 出发，以 45 度方向向右下移动；到达屏幕边界后根据碰撞位置改变方向并继续运动。程序总大小不得超过 510 字节。")
    add_body(doc, "[此处补充字符位置更新、边界判断、延时控制、颜色变化或其他扩展效果的设计过程。]", color=MUTED)


def add_code(doc):
    add_h1(doc, "3. 关键代码")
    add_body(doc, "本节暂不写入具体代码。完成实验后，请仅粘贴能够体现核心思路的片段，并在代码块下方解释寄存器含义、控制流程及其与实验结果之间的对应关系。")
    add_code_placeholder(doc, "3.1 Assignment 1：指定位置与字符属性设置",
                         "建议粘贴：坐标换算、显存写入或等效输出逻辑。")
    add_code_placeholder(doc, "3.2 Assignment 2：光标中断、字符输出与键盘回显",
                         "建议粘贴：光标位置获取/移动、中断输出学号、读取键盘并回显的核心片段。")
    add_code_placeholder(doc, "3.3 Assignment 3：分支逻辑、循环逻辑与 your_function",
                         "建议分为三个小片段粘贴，并逐段说明 32 位寄存器和栈操作。")
    add_code_placeholder(doc, "3.4 Assignment 4：字符移动与边界反弹",
                         "建议粘贴：位置更新、方向切换、边界判断和可选扩展效果。")


def add_results(doc):
    add_h1(doc, "4. 实验结果")
    add_body(doc, "请使用本人实际运行截图替换下方占位框，并在每张图片后补充一至两句话，说明截图展示的现象、是否符合预期，以及它验证了哪一项实验要求。")

    add_h2(doc, "4.1 Assignment 1：MBR")
    add_figure_placeholder(doc, "图 1  Example 1 复现结果", "粘贴 Example 1 正常启动和输出的截图")
    add_body(doc, "[补充说明：截图中的输出内容及其验证意义。]", color=MUTED)
    add_figure_placeholder(doc, "图 2  在 (12,12) 处显示学号的结果", "粘贴指定坐标、指定颜色显示本人学号的截图")
    add_body(doc, "[补充说明：坐标、前景色、背景色与预期是否一致。]", color=MUTED)

    add_h2(doc, "4.2 Assignment 2：实模式中断")
    add_figure_placeholder(doc, "图 3  光标位置获取与移动结果", "粘贴光标移动前后或光标位置验证截图")
    add_body(doc, "[补充说明：光标最终位置，以及使用的中断功能如何得到验证。]", color=MUTED)
    add_figure_placeholder(doc, "图 4  使用中断输出学号的结果", "粘贴通过实模式中断输出本人学号的截图")
    add_body(doc, "[补充说明：输出内容及其与 Assignment 1 实现方式的区别。]", color=MUTED)
    add_figure_placeholder(doc, "图 5  键盘输入与回显结果", "粘贴键盘输入并回显的截图")
    add_body(doc, "[补充说明：输入内容、回显内容和结束方式。]", color=MUTED)

    add_h2(doc, "4.3 Assignment 3：汇编逻辑")
    add_figure_placeholder(doc, "图 6  make run 测试结果", "粘贴 assignment 目录下执行 make run 的终端截图")
    add_body(doc, "[补充说明：if test、while test 和字符串输出是否均符合预期；如修改过 a1、a2，请记录测试值。]", color=MUTED)

    add_h2(doc, "4.4 Assignment 4：字符弹射程序")
    add_figure_placeholder(doc, "图 7  字符弹射程序运行结果", "粘贴字符运动、边界反弹或扩展效果的截图")
    add_body(doc, "[补充说明：字符初始位置、反弹路径、程序大小，以及本人加入的可选效果。]", color=MUTED)


def add_summary(doc):
    add_h1(doc, "5. 总结")
    add_body(doc, "本实验从 MBR 的加载与显示入手，逐步扩展到实模式中断、32 位汇编控制逻辑和字符弹射程序。通过四项任务，可以建立对 BIOS 加载 MBR、实模式程序执行、寄存器使用、内存寻址、控制转移、过程调用和中断服务的整体认识。")
    add_body(doc, "完成实验后，请结合本人真实经历补充以下内容：")
    add_bullet(doc, "最耗时的问题是什么，如何定位到原因。")
    add_bullet(doc, "MBR 直接操作显存与通过 BIOS 中断输出字符的差异。")
    add_bullet(doc, "Assignment 3 中 16 位与 32 位寄存器使用要求的区别。")
    add_bullet(doc, "字符弹射程序如何在 510 字节限制内完成位置更新、边界判断和显示。")
    add_body(doc, "[此处补充本人总结、排错过程和可改进之处。]", color=MUTED)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(12)
    for style_name in ("List Number", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(12)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("操作系统原理实验报告  |  实验二"), size=9, color=MUTED)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    set_styles(doc)
    add_cover(doc)
    add_requirements(doc)
    add_process(doc)
    add_code(doc)
    add_results(doc)
    add_summary(doc)
    add_footer(doc)
    doc.core_properties.title = "操作系统实验二报告填写版"
    doc.core_properties.subject = "实模式和保护模式下的 OS 启动"
    doc.core_properties.author = ""
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
